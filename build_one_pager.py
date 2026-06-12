"""Generate the MailSort one-pager as a .docx file.

Run:  python build_one_pager.py
Output: MailSort-OnePager-BestEffortRejectRecovery.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\work\ELC\AI Detection\MailSort-OnePager-BestEffortRejectRecovery.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREY_FILL = "F2F2F2"

doc = Document()

# Page margins (one-pager should be tight).
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# Default body style.
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(text, level=1, color=NAVY, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if size is None:
        size = {1: 18, 2: 13, 3: 11}.get(level, 11)
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_para(text, bold=False, italic=False, color=None, size=None, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    if not p.runs:
        run = p.add_run(text)
    else:
        run = p.runs[0]
        run.text = text
    return p


def add_kv_bullet(label, value, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label)
    r1.bold = True
    p.add_run(" " + value)
    return p


# ---------- Title ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
t_run = title.add_run("Best-Effort Reject-Bin Recovery with Offline AI-Assisted Handwriting")
t_run.bold = True
t_run.font.size = Pt(18)
t_run.font.color.rgb = NAVY
title.paragraph_format.space_after = Pt(2)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
s_run = subtitle.add_run(
    "Proposed by: MailSort Engineering    |    Status: Proposal for review    |    "
    "Target hardware: ASUS GX10 (NVIDIA GB10, 128 GB unified memory) - dedicated, on-premises"
)
s_run.italic = True
s_run.font.size = Pt(10)
s_run.font.color.rgb = GREY
subtitle.paragraph_format.space_after = Pt(6)

# ---------- Business problem ----------
add_heading("Business problem", level=2)
add_para(
    "When an envelope's 2D barcode is damaged, missing, or unreadable, the sorter physically diverts "
    "the envelope to a reject tray. Today, every rejected envelope becomes manual work: an operator "
    "retrieves the bin, visually identifies each envelope, and re-feeds it through the sorter after "
    "typing in the recipient data. This is the slowest, most expensive, and least consistent part of "
    "the workflow, and it scales linearly with bad-barcode volume."
)
add_para(
    "We need a way to fully automate the recovery of the majority of those envelopes - no operator "
    "touch, no manual data entry - using on-premises hardware, without ever replacing the barcode as "
    "the primary signal, and without slowing the sorter."
)

# ---------- Proposed solution ----------
add_heading("Proposed solution", level=2)
add_para(
    "We give every bad-barcode envelope a persistent identity based on a stable visual fingerprint of "
    "the envelope image. We then decouple the sorter's real-time decision from the slow work of "
    "reading the handwriting, joining the two via a single database record. Every actor in the system - "
    "the sorter, the AI server, the operator workstation - reads and writes the same record, keyed by "
    "the fingerprint."
)
add_para(
    "The system has two completely separate time regimes that communicate only through the database. "
    "They never wait for each other."
)

add_heading("Regime 1 - the sorter, hard real-time (<= 400 ms per envelope)", level=3)
add_para("On every scan, the sorter hands the application the envelope image. The application:")
add_bullet("Computes a stable envelope fingerprint from the image (under 100 ms).")
add_bullet(
    "Looks up that fingerprint in a database table. If a verified recipient record exists for it, "
    "the application tells the sorter the correct tray in under 5 ms. The envelope is routed."
)
add_bullet(
    "If no verified record exists, the application tells the sorter to send the envelope to the "
    "reject tray, persists a row with the fingerprint, and returns. Total time: under 100 ms typical, "
    "well under the 400 ms budget."
)
add_para(
    "The sorter app never calls the AI server, never blocks on a network call, never reads or "
    "interprets handwriting. It is a fingerprint in, tray out. This guarantees the sorter runs at "
    "full speed no matter what is happening elsewhere.",
    italic=True,
    color=GREY,
)

add_heading("Regime 2 - the back office, soft real-time (minutes to hours)", level=3)
add_para(
    "Out-of-band, with no time pressure, a worker process takes the saved image, crops the name and "
    "ID regions (using the same per-field ROIs the sorter app already configures), and sends the two "
    "crops to the dedicated on-premises AI server (ASUS GX10, NVIDIA GB10). The AI server returns the "
    "handwritten name and ID. The worker then checks our recipient registry: is this (name, ID) pair "
    "a real record?"
)
add_kv_bullet(
    "Yes, with high confidence on both fields ->",
    "the worker writes verified = true, the recipient's tray, and a timestamp to the envelope's row, "
    "marked verified_by = 'ai'. The envelope is fully automated - no operator ever touches it. The "
    "next sorter run that sees this fingerprint routes it correctly on its very first lookup.",
)
add_kv_bullet(
    "AI reads the fields but the (name, ID) pair is not in the registry, or the AI's confidence is "
    "below threshold ->",
    "the worker notifies an operator workstation. The operator sees the AI's best guess, confirms or "
    "corrects it, and saves. Operator time is reduced because they are confirming, not starting from "
    "scratch.",
)
add_kv_bullet(
    "After N reject cycles with no verification (default: 3) ->",
    "the envelope is permanently routed to a manual-entry workstation queue. A human handles it "
    "directly, and their entry is what unblocks the next sorter run.",
)

add_heading("Why this design works", level=3)
add_bullet("The sorter cannot be slowed down by the AI server. The 400 ms decision routine never calls the AI server. It is a fingerprint in, tray out, full stop.")
add_bullet(
    "The AI server cannot take down the sorter. If the AI server is offline, slow, or unsure, the "
    "worst case is \"every bad-barcode envelope is rejected, just like today.\" Nothing we add makes "
    "a working system worse."
)
add_bullet(
    "Privacy and compliance are preserved. Every envelope image and every OCR result stays on our "
    "network. The recipient registry is never queried by a third party. There is no per-envelope "
    "cloud spend and no envelope image leaves the building."
)
add_bullet(
    "The system is self-improving. Every envelope processed produces an audit row: envelope_id, "
    "ai_name, ai_id, ai_confidence, manual_name, manual_id, agreed. That audit table is the training "
    "data for the next quarterly model fine-tune. The more mail we process, the better the AI gets, "
    "and the higher the fully-automated rate climbs."
)

# ---------- Target outcome ----------
add_heading("Target outcome and how we measure it", level=2)
add_para(
    "The single number that matters is the fully-automated rate: the percentage of bad-barcode "
    "envelopes that are routed end-to-end with zero operator touch. A \"fully automated\" envelope is "
    "one where:"
)
add_bullet("The AI read the name with confidence above threshold, and")
add_bullet("The AI read the ID with confidence above threshold, and")
add_bullet("The (name, ID) pair was found in the recipient registry, and")
add_bullet(
    "A verified = 1, verified_by = 'ai' row exists before the next sorter run sees the envelope."
)

pilot = doc.add_paragraph()
pilot_run = pilot.add_run("Pilot commitment:")
pilot_run.bold = True
pilot.add_run(
    "  the 30-day pilot must demonstrate a fully-automated rate of at least 60% (floor) and 80% "
    "(stretch target) on real mail volume. Sub-rates reported every week: OCR name confidence pass "
    "rate, OCR ID confidence pass rate, registry match rate on OCR-passed items, and the headline "
    "fully-automated rate. The sub-rates tell us where to invest if the floor is missed."
)
pilot.paragraph_format.space_after = Pt(4)

add_para(
    "A fully-automated rate below 60% does not mean \"abandon the project.\" It means one of the four "
    "sub-rates is the bottleneck, and the pilot report identifies which one - and the fix is targeted "
    "(a model upgrade, a registry cleanup, a ROI adjustment, or a workstation UX change), not a redo.",
    italic=True,
)

# ---------- Secondary value ----------
add_heading("Secondary value - the \"failed\" path is still faster", level=2)
add_para(
    "For envelopes the AI cannot fully automate, the operator sees only the cropped name and ID "
    "regions, not the whole envelope, and confirms or corrects the AI's best guess - they never start "
    "from a blank page. Operator time is reduced even on the \"failed\" path; the AI's failure does "
    "not waste the human's time, it only adds a confirmation step. We expect the operator's per-"
    "envelope time on AI-flagged items to be roughly half of today's manual-entry time, even on "
    "envelopes the AI could not fully route."
)

# ---------- What this is not ----------
add_heading("What this is not", level=2)
add_bullet("Not a replacement for the barcode. The barcode remains the primary signal. AI-OCR is a fallback for the reject bin.")
add_bullet(
    "Not a one-shot automation push. The AI model is fine-tuned quarterly on our own envelope crops, "
    "and the fully-automated rate is re-measured every quarter. The number goes up over time; this "
    "proposal buys the data flywheel, not just the first deployment."
)
add_bullet("Not a cloud service. Every part of this proposal runs on hardware we own, in our facility, on our network. There is no per-envelope cloud spend and no envelope image leaves the building.")
add_bullet("Not coupled to the sorter's real-time loop. The AI server is a back-office system. The sorter runs at full speed regardless of its state.")

# ---------- Cost shape ----------
add_heading("Cost shape", level=2)
cost_rows = [
    ("Item", "Direction", "Notes"),
    ("ASUS GX10 (NVIDIA GB10) appliance", "One-time capital", "Dedicated, rack-mountable, no specialized power or cooling"),
    ("Handwriting OCR model (e.g. TrOCR-base)", "One-time, free", "Open source; optional fine-tune on our own envelope images"),
    (
        "Engineering: envelope-fingerprint join key, DB schema, sorter hot path, AI worker, manual-entry UI",
        "Internal time",
        "~1-2 sprints, builds on the existing MailSort pHash pipeline",
    ),
    (
        "Engineering: pilot telemetry + 30-day measurement",
        "Internal time",
        "Same work that pays for the pilot also instruments the production system",
    ),
    (
        "Cloud OCR API (comparison arm of pilot only)",
        "Pay-per-use, time-boxed",
        "Optional. Allows head-to-head accuracy comparison against the local model",
    ),
    (
        "Recurring: quarterly model fine-tune on collected envelope crops",
        "Internal time",
        "The audit log is the training set",
    ),
]
table = doc.add_table(rows=len(cost_rows), cols=3)
table.autofit = True
for i, (a, b, c) in enumerate(cost_rows):
    row = table.rows[i]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    for j, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if i == 0:
            set_cell_shading(cell, "1F3A5F")
        else:
            set_cell_shading(cell, LIGHT_GREY_FILL if i % 2 == 0 else "FFFFFF")

add_para(
    "The AI server is sized for worst-case mail volume plus headroom. A single GX10 is sufficient; "
    "the pilot will confirm.",
    italic=True,
    color=GREY,
)

# ---------- Risks ----------
add_heading("Risks and how we mitigate them", level=2)
risks = [
    (
        "Fully-automated rate is below the 60% floor",
        "Pilot reports sub-rates weekly so the bottleneck is identified early. Each sub-rate has a known, targeted fix path.",
    ),
    (
        "AI model is less accurate than expected on our real envelopes",
        "Pilot measures on 10,000+ real items before commitment. Optional head-to-head against a cloud OCR API. If neither hits the floor, we stop and have lost only the pilot cost.",
    ),
    (
        "AI server is a new single point of failure",
        "Designed-in: if the AI server is unreachable, every bad-barcode envelope is rejected, exactly like today. No regression.",
    ),
    (
        "AI server is too slow to keep up with mail volume",
        "GB10 at ~1 PFLOPS has massive headroom for OCR. The 30-day pilot measures throughput under real load.",
    ),
    (
        "Fingerprint collisions (two different envelopes hash to the same key)",
        "Pilot instruments and reports collision rate. If non-zero, we add a stronger perceptual hash or a secondary discriminator. Bounded by DB unique constraint.",
    ),
    (
        "Privacy / compliance concern about an AI reading mail",
        "100% on-premises. No outbound network calls from the AI server. Envelope images never leave the facility.",
    ),
    (
        "Staff resist the change",
        "AI-OCR is positioned as a productivity tool that removes the most tedious part of their day, not a replacement. The proposal explicitly protects their role for the cases the AI cannot read, and the AI's best guess is presented as a suggestion, not a fait accompli.",
    ),
]
for label, value in risks:
    p = doc.add_paragraph()
    r1 = p.add_run(label + " - ")
    r1.bold = True
    p.add_run(value)
    p.paragraph_format.space_after = Pt(2)

# ---------- Ask ----------
add_heading("Ask", level=2)
add_bullet("Approval to run a 30-day pilot on a single sorter line, with the GX10 appliance leased or purchased for the duration.")
add_bullet("Approval to commit a 60% fully-automated floor as the deployment threshold, measured on real mail volume.")
add_bullet("A named executive sponsor to receive weekly pilot reports and the go/no-go recommendation at day 30.")
add_para(
    "If the pilot meets the bar, the rollout to the remaining sorter lines is a configuration change, "
    "not a new project. If it misses the bar, the pilot report tells us exactly which of the four sub-"
    "rates to invest in next, and the data flywheel continues regardless.",
    italic=True,
)

# ---------- Appendix ----------
doc.add_page_break()
add_heading("Appendix - Technical detail (for engineering review)", level=2)
add_para(
    "This appendix is intended for the engineering reviewer and is not part of the one-page narrative.",
    italic=True,
    color=GREY,
)

add_heading("Pipeline integration with the existing MailSort code base", level=3)
add_bullet(
    "The existing RegionalFingerprint already produces three perceptual hash channels. We add two "
    "more ROIs, NameRoi and IdRoi, configured in MatchSettings the same way as AddressRoi/BarcodeRoi. "
    "The same crop + deskew + contrast-stretch pipeline produces both the join-key hash and the OCR "
    "inputs - no duplicated preprocessing."
)
add_bullet(
    "Join key: the fingerprint used for the reject-bin join must be a deterministic exact-match key "
    "(an indexed 64-bit or 128-bit hash), not a similarity search. The cleanest production approach "
    "is to promote the most stable of the existing pHash channels to a UNIQUE indexed column "
    "(envelope_lookup.envelope_id) and keep the others as fallback fuzzy-lookup fields. This is a "
    "small, well-scoped change to the persistence layer."
)
add_bullet(
    "IngestService.IngestAsync already runs in a synchronous, sub-second loop called from the sorter "
    "endpoint. We add a fast path: compute the join key, do an indexed SELECT tray FROM "
    "envelope_lookup WHERE envelope_id = ? AND verified = 1, return tray. If no row, persist a new "
    "verified = 0 row, fire-and-forget the AI worker job, return the reject tray."
)

add_heading("Hot-path latency budget (the 400 ms wall)", level=3)
add_bullet("0-50 ms - decode JPEG, compute envelope_id (join key)")
add_bullet("50-60 ms - DB lookup (indexed)")
add_bullet("60-100 ms - return tray")
add_bullet("Headroom: 300 ms")
add_bullet("The 400 ms budget is never approached; the GB10 is never called.")

add_heading("Back-office worker", level=3)
add_bullet(
    "Runs on a separate process / host. It reads the queue of (envelope_id, image_path, "
    "roi_name_path, roi_id_path) jobs, calls the GB10 over LAN gRPC, and writes results back. It "
    "uses the same MailSortDbContext for the recipient registry lookup."
)
add_bullet(
    "Recipient registry lookup: parameterized SELECT id, tray FROM recipients WHERE "
    "normalized_name = ? AND id_number = ? with a composite index on (normalized_name, id_number). "
    "On match, the worker's UPDATE envelope_lookup SET verified = 1, recipient_id = ?, tray = ?, "
    "verified_at = UTCNOW(), verified_by = 'ai' unblocks the next sorter run."
)
add_bullet(
    "Manual-entry UI: a small Blazor page in MailSort/Components/Pages that reads envelope_lookup "
    "rows where verified = 0 AND notified_at IS NOT NULL, displays the cropped name and ID ROIs "
    "along with the AI's best guess, accepts typed input, and writes back. No new framework."
)
add_bullet(
    "Retry bound: a nightly job (or per-scan trigger) counts reject cycles per envelope. When the "
    "count exceeds the configured N, the row is routed to the manual-entry workstation queue "
    "permanently."
)
add_bullet(
    "Audit / training data table (envelope_ai_audit): (envelope_id, ai_name, ai_id, ai_confidence, "
    "ai_latency_ms, manual_name, manual_id, agreed, created_at). This is the dataset used to "
    "fine-tune the OCR model on our own envelope crops. We do not need a separate ML platform; the "
    "existing MailSortDbContext handles it."
)

add_heading("Recommended model on the GB10", level=3)
add_bullet(
    "TrOCR-base handwritten (~330 MB, ~50-80 ms per crop on GB10) as the default. Fine-tune on "
    "500-2000 of our own envelope crops for a 5-15% accuracy lift."
)
add_bullet(
    "Defer larger VLMs (Qwen2.5-VL-3B/7B) to a batch re-scan job that runs at end of shift. They "
    "are accurate but too slow for the hot path on any hardware."
)
add_bullet(
    "Serve via Triton + ONNX Runtime (lower overhead than vLLM for small batch sizes) or vLLM if "
    "we go generative. Expose one HTTP endpoint per field type (/ocr/name, /ocr/id). Pin the model "
    "warm in VRAM with --max-model-len 64."
)

add_heading("Sub-rate instrumentation (what the pilot dashboard must show)", level=3)
sub_rows = [
    ("Sub-rate", "Definition", "Pilot target"),
    ("OCR name confidence pass rate", "% of envelopes where the AI's name confidence >= threshold", ">= 85%"),
    ("OCR ID confidence pass rate", "% of envelopes where the AI's ID confidence >= threshold", ">= 90%"),
    ("Registry match rate on OCR-passed items", "% of OCR-passed envelopes where (name, ID) exists in the registry", ">= 90%"),
    ("Fully-automated rate", "% of bad-barcode envelopes with all three above and verified_by = 'ai' before next scan", ">= 60% floor, 80% target"),
]
t2 = doc.add_table(rows=len(sub_rows), cols=3)
for i, (a, b, c) in enumerate(sub_rows):
    row = t2.rows[i]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if i == 0:
            set_cell_shading(cell, "1F3A5F")
        else:
            set_cell_shading(cell, LIGHT_GREY_FILL if i % 2 == 0 else "FFFFFF")

add_para(
    "The product of the three sub-rates is the fully-automated rate. If the floor is missed, the "
    "lowest sub-rate is the bottleneck, and the fix path is known:",
    italic=True,
    color=GREY,
)
add_bullet("Lowest is OCR name -> upgrade or fine-tune the model, or relax the confidence threshold.")
add_bullet("Lowest is OCR ID -> same, or adjust the ID ROI.")
add_bullet("Lowest is registry match -> registry cleanup, or add the missing recipient.")
add_bullet("Lowest is end-to-end timing -> scale the worker pool or move to a faster model.")

add_heading("Failure-mode behavior, in priority order", level=3)
add_bullet("AI server down -> all flagged envelopes are rejected, exactly like today. Sorter keeps running.")
add_bullet("AI server slow -> no effect on sorter; back-office worker queue grows. Manual-entry UI still works.")
add_bullet("AI returns a (name, ID) that does not exist in the registry -> row stays verified = 0, operator is notified. We never auto-route on uncertain data.")
add_bullet("AI returns high confidence but the human disagrees on audit -> recorded, used as fine-tune data next quarter.")
add_bullet("Fingerprint collision -> bounded by DB UNIQUE constraint. If it ever happens, the envelope goes to manual entry; we add a secondary discriminator to the join key.")

doc.save(OUT)
print("Wrote:", OUT)
